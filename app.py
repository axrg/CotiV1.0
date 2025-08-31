import os
import uuid
from flask import Flask, render_template_string, request, send_from_directory
from docx import Document

app = Flask(__name__)

# Carpeta temporal para guardar los archivos
TMP_FOLDER = "static/tmp"
os.makedirs(TMP_FOLDER, exist_ok=True)

# HTML responsivo
FORM_HTML = """
<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Generar Cotización</title>
<style>
    body { font-family: Arial; background: #f7f7f7; margin: 0; padding: 0; }
    .container { max-width: 900px; width: 95%; margin: 20px auto; background: #fff; padding: 20px; border-radius: 10px; box-shadow: 0px 0px 15px rgba(0,0,0,0.2); }
    h2, h3 { text-align: center; color: #333; }
    form { display: flex; flex-direction: column; gap: 15px; }
    label { font-weight: bold; margin-bottom: 5px; }
    input[type="text"], input[type="number"], input[type="date"], select, button { padding: 10px; border-radius: 5px; border: 1px solid #ccc; width: 100%; box-sizing: border-box; }
    button { background-color: #007bff; color: #fff; font-size: 16px; cursor: pointer; border: none; }
    button:hover { background-color: #0056b3; }
    .concepto-item { display: flex; gap: 5px; flex-wrap: wrap; align-items: center; }
    .concepto-item input, .concepto-item select { flex: 1 1 120px; }
    #conceptos-container { margin-bottom: 10px; }
    .add-btn { margin-bottom: 20px; }
    .totales { font-weight: bold; text-align: right; margin-top: 10px; }
    @media(max-width:600px){
        .concepto-item { flex-direction: column; }
        .concepto-item input, .concepto-item select, .concepto-item button { width: 100%; }
    }
</style>
<script>
function agregarConcepto() {
    const container = document.getElementById("conceptos-container");
    const div = document.createElement("div");
    div.className = "concepto-item";
    div.innerHTML = `
        <input type="text" name="concepto[]" placeholder="Concepto" required>
        <input type="number" name="cantidad[]" placeholder="Cantidad" min="0" step="0.01" required oninput="calcularTotales()">
        <select name="unidad[]">
            <option value="Pza">Pza</option>
            <option value="Kg">Kg</option>
            <option value="m">m</option>
            <option value="L">L</option>
        </select>
        <input type="number" name="valor_unitario[]" placeholder="Valor Unitario" min="0" step="0.01" required oninput="calcularTotales()">
        <input type="text" class="subtotal" placeholder="Subtotal" readonly>
        <button type="button" onclick="this.parentElement.remove(); calcularTotales()">Eliminar</button>
    `;
    container.appendChild(div);
    calcularTotales();
}

function calcularTotales() {
    let totalMateriales = 0;
    const items = document.querySelectorAll(".concepto-item");
    items.forEach(item => {
        const cantidad = parseFloat(item.querySelector('input[name="cantidad[]"]').value) || 0;
        const valor = parseFloat(item.querySelector('input[name="valor_unitario[]"]').value) || 0;
        const subtotal = cantidad * valor;
        item.querySelector('.subtotal').value = `$${subtotal.toFixed(2)} MXN`;
        totalMateriales += subtotal;
    });
    const manoObra = parseFloat(document.getElementById("mano_obra").value) || 0;
    const gestion = parseFloat(document.getElementById("gestion").value) || 0;
    const totalGeneral = totalMateriales + manoObra + gestion;
    document.getElementById("total_materiales").innerText = `$${totalMateriales.toFixed(2)} MXN`;
    document.getElementById("total_general").innerText = `$${totalGeneral.toFixed(2)} MXN`;
}

window.onload = function() {
    agregarConcepto();
    document.getElementById("mano_obra").oninput = calcularTotales;
    document.getElementById("gestion").oninput = calcularTotales;
};
</script>
</head>
<body>
<div class="container">
    <h2>Generar Cotización</h2>
    <form method="POST" action="/generar">
        <label>Fecha:</label>
        <input name="fecha" type="date" required>
        <label>Nombre del cliente:</label>
        <input name="nombre_cliente" type="text" placeholder="Ej: Juan Pérez" required>
        <label>Título de cotización:</label>
        <input name="titulo_cotizacion" type="text" placeholder="Ej: Instalación de calentadores">
        <label>Plazo de la oferta:</label>
        <input name="plazo_oferta" type="text" placeholder="Ej: 15 días">
        <label>Tiempo de entrega:</label>
        <input name="tiempo_entrega" type="text" placeholder="Ej: 7 días hábiles">
        <label>Pago acordado:</label>
        <input name="pago_acordado" type="text" placeholder="Ej: 50% anticipo">

        <h3>Conceptos</h3>
        <div id="conceptos-container"></div>
        <button type="button" class="add-btn" onclick="agregarConcepto()">Agregar concepto</button>

        <label>Mano de obra:</label>
        <input name="mano_obra" id="mano_obra" type="number" value="500" min="0" step="0.01">
        <label>Gestión:</label>
        <input name="gestion" id="gestion" type="number" value="100" min="0" step="0.01">

        <div class="totales">
            Total materiales: <span id="total_materiales">$0.00 MXN</span> | Total general: <span id="total_general">$0.00 MXN</span>
        </div>

        <button type="submit">Generar Word</button>
    </form>
</div>
</body>
</html>
"""

@app.route("/")
def index():
    return render_template_string(FORM_HTML)

@app.route("/generar", methods=["POST"])
def generar():
    datos_generales = {
        "fecha": request.form["fecha"],
        "nombre_cliente": request.form["nombre_cliente"],
        "titulo_cotizacion": request.form["titulo_cotizacion"],
        "plazo_oferta": request.form["plazo_oferta"],
        "tiempo_entrega": request.form["tiempo_entrega"],
        "pago_acordado": request.form["pago_acordado"]
    }

    # Conceptos
    conceptos = []
    total_materiales = 0
    for i, (c, cant, u, v) in enumerate(zip(
        request.form.getlist("concepto[]"),
        request.form.getlist("cantidad[]"),
        request.form.getlist("unidad[]"),
        request.form.getlist("valor_unitario[]")
    ), start=1):
        cantidad = float(cant)
        valor_unitario = float(v)
        subtotal = cantidad * valor_unitario
        total_materiales += subtotal
        conceptos.append({
            "indice": i,
            "concepto": c,
            "cantidad": cantidad,
            "unidad": u,
            "valor_unitario": valor_unitario,
            "subtotal": subtotal
        })

    mano_obra = float(request.form.get("mano_obra", 0))
    gestion = float(request.form.get("gestion", 0))
    total_general = total_materiales + mano_obra + gestion

    # Crear docx
    doc = Document("plantilla1.docx")
    for para in doc.paragraphs:
        for key, value in datos_generales.items():
            para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
        para.text = para.text.replace("${{total_materiales}}", f"${total_materiales:.2f} MXN")
        para.text = para.text.replace("${{mano_obra}}", f"${mano_obra:.2f} MXN")
        para.text = para.text.replace("${{gestion}}", f"${gestion:.2f} MXN")
        para.text = para.text.replace("${{total_general}}", f"${total_general:.2f} MXN")

    # Tabla conceptos
    tabla = doc.tables[0]
    fila_ejemplo = tabla.rows[1]
    tabla._tbl.remove(fila_ejemplo._tr)
    for c in conceptos:
        fila = tabla.add_row()
        fila.cells[0].text = str(c["indice"])
        fila.cells[1].text = c["concepto"]
        fila.cells[2].text = f"{c['cantidad']:.2f}"
        fila.cells[3].text = c["unidad"]
        fila.cells[4].text = f"${c['valor_unitario']:.2f} MXN"
        fila.cells[5].text = f"${c['subtotal']:.2f} MXN"

    # Tabla resumen
    tabla_resumen = doc.tables[1]
    while len(tabla_resumen.rows) < 4:
        tabla_resumen.add_row()
    tabla_resumen.cell(0,1).text = f"${total_materiales:.2f} MXN"
    tabla_resumen.cell(1,1).text = f"${mano_obra:.2f} MXN"
    tabla_resumen.cell(2,1).text = f"${gestion:.2f} MXN"
    tabla_resumen.cell(3,1).text = f"${total_general:.2f} MXN"

    # Guardar docx temporal
    filename = f"cotizacion_{uuid.uuid4().hex}.docx"
    filepath = os.path.join(TMP_FOLDER, filename)
    doc.save(filepath)

    return f"""
    <h3>Tu cotización está lista ✅</h3>
    <p>Descárgala aquí: <a href="/descargar/{filename}">Descargar Word</a></p>
    """

@app.route("/descargar/<filename>")
def descargar(filename):
    return send_from_directory(TMP_FOLDER, filename, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)


